"""
Advanced TOC and Style Extractor for Proposal Templates

This module extracts comprehensive template information from DOCX files:
1. Table of Contents structure with proper hierarchy
2. Section word count patterns and ranges
3. Writing style, tone, and content organization
4. Information flow and presentation patterns
"""

import os
import json
import re
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path

try:
    from docx import Document
    from docx.document import Document as DocumentType
    from docx.text.paragraph import Paragraph
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx not available, TOC extraction will be limited")


class TemplateAnalyzer:
    """Analyzes DOCX files to extract comprehensive template information"""

    def __init__(self):
        self.templates_dir = Path("toc_templates")
        self.templates_dir.mkdir(exist_ok=True)

    def extract_text_from_paragraph(self, paragraph) -> str:
        """Extract text from paragraph, handling different formatting"""
        if not paragraph:
            return ""
        return paragraph.text.strip()

    def analyze_paragraph_style(self, paragraph) -> Dict[str, Any]:
        """Analyze paragraph formatting and style - AGGRESSIVE heading detection"""
        text = self.extract_text_from_paragraph(paragraph)
        style_info = {
            'text': text,
            'style_name': paragraph.style.name if paragraph.style else 'Normal',
            'alignment': str(paragraph.alignment) if paragraph.alignment else 'LEFT',
            'is_bold': False,
            'is_italic': False,
            'font_size': None,
            'is_heading': False,
            'heading_level': 0
        }

        # Check formatting in runs first
        max_font_size = 0
        all_bold = True
        has_text = False

        for run in paragraph.runs:
            if run.text.strip():  # Only consider runs with actual text
                has_text = True
                if run.bold:
                    style_info['is_bold'] = True
                else:
                    all_bold = False
                if run.italic:
                    style_info['is_italic'] = True
                if run.font.size:
                    font_size = run.font.size.pt
                    style_info['font_size'] = font_size
                    max_font_size = max(max_font_size, font_size)

        # If no text runs, don't consider it all bold
        if not has_text:
            all_bold = False

        # PRIORITY 1: Actual Word heading styles
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()

            # Check for numbered heading styles (common variations)
            if style_name == 'heading 1':
                style_info['is_heading'] = True
                style_info['heading_level'] = 1
            elif style_name == 'heading 2':
                style_info['is_heading'] = True
                style_info['heading_level'] = 2
            elif style_name == 'heading 3':
                style_info['is_heading'] = True
                style_info['heading_level'] = 3
            elif style_name == 'heading 4':
                style_info['is_heading'] = True
                style_info['heading_level'] = 4
            elif style_name == 'heading 5':
                style_info['is_heading'] = True
                style_info['heading_level'] = 5
            elif style_name == 'heading 6':
                style_info['is_heading'] = True
                style_info['heading_level'] = 6
            elif style_name in ['title', 'subtitle']:
                style_info['is_heading'] = True
                style_info['heading_level'] = 1

        # PRIORITY 2: If not a heading style, check formatting patterns
        if not style_info['is_heading'] and text.strip():
            # Pattern 1: Short text (< 150 chars), all bold, larger font OR specific font sizes
            if (len(text.strip()) < 150 and
                (all_bold and max_font_size >= 12) or
                (max_font_size >= 14)):  # Bradford uses 14pt for Heading 1, 12pt for Heading 2

                style_info['is_heading'] = True
                # Determine level based on font size (Bradford-specific)
                if max_font_size >= 14:
                    style_info['heading_level'] = 1  # 14pt = Heading 1
                elif max_font_size >= 12:
                    style_info['heading_level'] = 2  # 12pt = Heading 2
                elif max_font_size >= 10:
                    style_info['heading_level'] = 3  # 10pt = Heading 3
                else:
                    style_info['heading_level'] = 4

                print(f"🔍 DETECTED heading by formatting: '{text.strip()}' (Size: {max_font_size}, Bold: {all_bold})")

            # Pattern 2: Text that looks like a section header (numbered, etc.)
            elif len(text.strip()) < 100:
                text_lower = text.strip().lower()
                # Check for common section patterns
                heading_patterns = [
                    r'^\d+\.?\s',  # "1. " or "1 "
                    r'^[a-z]\)?\s',  # "a) " or "a "
                    r'^\w+\s*:$',  # "Introduction:"
                    r'^(executive|technical|commercial|financial|project|scope|timeline|budget|team|approach|methodology|deliverables|conclusion|appendix)',
                    r'^(table of contents|introduction|overview|summary|background|requirements|solution|implementation|testing|deployment|maintenance|support)'
                ]

                for pattern in heading_patterns:
                    import re
                    if re.search(pattern, text_lower):
                        style_info['is_heading'] = True
                        # Default to level 2 for pattern-detected headers
                        style_info['heading_level'] = 2 if not all_bold else 1
                        print(f"🔍 DETECTED heading by pattern: '{text.strip()}' (Pattern: {pattern})")
                        break

            # Pattern 3: All caps text (likely headers)
            elif (len(text.strip()) < 80 and
                  text.strip().isupper() and
                  len(text.strip().split()) <= 6):
                style_info['is_heading'] = True
                style_info['heading_level'] = 2
                print(f"🔍 DETECTED heading by caps: '{text.strip()}'")

        return style_info

    def extract_toc_structure(self, doc: DocumentType) -> List[Dict[str, Any]]:
        """SIMPLE TOC extraction - just find ALL headings and organize them"""

        print(f"🔍 Starting SIMPLE TOC extraction from document with {len(doc.paragraphs)} paragraphs...")

        # Step 1: Find ALL headings in order
        all_headings = []

        for i, paragraph in enumerate(doc.paragraphs):
            style_info = self.analyze_paragraph_style(paragraph)
            text = style_info['text']

            if not text.strip():
                continue

            # Is this a heading?
            if style_info['is_heading'] and style_info['heading_level'] > 0:
                heading = {
                    'title': text,
                    'level': style_info['heading_level'],
                    'style': style_info,
                    'paragraph_index': i
                }
                all_headings.append(heading)
                print(f"📝 FOUND HEADING: '{text}' (Level {style_info['heading_level']})")

        print(f"✅ Found {len(all_headings)} total headings")

        # Step 2: Build hierarchy - simple approach
        sections = []
        current_main_section = None

        for heading in all_headings:
            if heading['level'] == 1:
                # This is a main section
                if current_main_section:
                    sections.append(current_main_section)

                current_main_section = {
                    'title': heading['title'],
                    'level': 1,
                    'style': heading['style'],
                    'subsections': [],
                    'content': [],
                    'word_count': 100  # Default word count
                }
                print(f"📁 MAIN SECTION: {heading['title']}")

            else:
                # This is a subsection (level 2, 3, etc.)
                if current_main_section:
                    subsection = {
                        'title': heading['title'],
                        'level': heading['level'],
                        'style': heading['style'],
                        'subsections': [],
                        'content': [],
                        'word_count': 50  # Default word count
                    }
                    current_main_section['subsections'].append(subsection)
                    print(f"  📄 SUBSECTION: {heading['title']} (Level {heading['level']})")
                else:
                    # No main section yet, treat as main section
                    current_main_section = {
                        'title': heading['title'],
                        'level': heading['level'],
                        'style': heading['style'],
                        'subsections': [],
                        'content': [],
                        'word_count': 100
                    }
                    print(f"📁 ORPHAN → MAIN SECTION: {heading['title']}")

        # Don't forget the last section
        if current_main_section:
            sections.append(current_main_section)

        print(f"✅ FINAL RESULT:")
        for i, section in enumerate(sections):
            subsection_count = len(section.get('subsections', []))
            print(f"  📋 Section {i+1}: '{section['title']}' ({subsection_count} subsections)")
            for j, subsection in enumerate(section.get('subsections', [])):
                print(f"    📄 {i+1}.{j+1}: '{subsection['title']}'")

        return sections

    def count_words(self, text: str) -> int:
        """Count words in text"""
        if not text:
            return 0
        return len(text.split())

    def analyze_writing_style(self, sections: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Analyze writing style patterns across sections - COMPREHENSIVE"""
        style_analysis = {
            'avg_sentence_length': 0,
            'avg_paragraph_length': 0,
            'common_phrases': [],
            'tone_indicators': [],
            'structural_patterns': [],
            'total_word_count': 0,
            'section_word_ranges': {},
            'writing_samples': {},  # Actual text samples from each section
            'table_usage_patterns': {},  # Where tables appear
            'formality_level': 'professional'  # formal, professional, casual
        }

        all_text = []
        total_words = 0
        writing_samples = {}
        table_patterns = {}

        for section in sections:
            # Get writing sample from this section
            section_sample = section.get('content_sample', '')
            section_title = section['title']
            section_words = section.get('word_count', 0)
            total_words += section_words

            # Store writing sample
            if section_sample:
                writing_samples[section_title] = section_sample
                all_text.append(section_sample)

            # Track table usage
            if section.get('has_tables'):
                table_patterns[section_title] = {
                    'count': section.get('table_count', 0),
                    'position': 'integrated',  # tables are part of content flow
                    'tables': section.get('tables', [])
                }

            # Track word count ranges by section title directly
            section_type = self.categorize_section(section_title.lower())

            if section_type not in style_analysis['section_word_ranges']:
                style_analysis['section_word_ranges'][section_type] = []
            style_analysis['section_word_ranges'][section_type].append(section_words)

            # Also process subsections
            for subsection in section.get('subsections', []):
                sub_sample = subsection.get('content_sample', '')
                sub_title = subsection['title']
                sub_words = subsection.get('word_count', 0)
                total_words += sub_words

                if sub_sample:
                    writing_samples[f"{section_title} > {sub_title}"] = sub_sample
                    all_text.append(sub_sample)

                if subsection.get('has_tables'):
                    table_patterns[f"{section_title} > {sub_title}"] = {
                        'count': subsection.get('table_count', 0),
                        'position': 'integrated',
                        'tables': subsection.get('tables', [])
                    }

                sub_type = self.categorize_section(sub_title.lower())
                if sub_type not in style_analysis['section_word_ranges']:
                    style_analysis['section_word_ranges'][sub_type] = []
                style_analysis['section_word_ranges'][sub_type].append(sub_words)

        style_analysis['writing_samples'] = writing_samples
        style_analysis['table_usage_patterns'] = table_patterns
        style_analysis['total_word_count'] = total_words

        full_text = ' '.join(all_text)

        # Analyze sentence patterns
        sentences = re.split(r'[.!?]+', full_text)
        sentences = [s.strip() for s in sentences if s.strip()]

        if sentences:
            avg_sentence_words = sum(len(s.split()) for s in sentences) / len(sentences)
            style_analysis['avg_sentence_length'] = round(avg_sentence_words, 1)

        # Find common business phrases
        business_phrases = [
            'we propose', 'our approach', 'we recommend', 'our solution',
            'we deliver', 'our team', 'we understand', 'our experience',
            'in conclusion', 'furthermore', 'however', 'therefore',
            'our commitment', 'we are pleased', 'we will', 'shall', 'will provide'
        ]

        found_phrases = []
        for phrase in business_phrases:
            if phrase in full_text.lower():
                found_phrases.append(phrase)

        style_analysis['common_phrases'] = found_phrases

        # Determine formality level
        formal_indicators = ['shall', 'hereby', 'wherein', 'whereas', 'aforementioned']
        formal_count = sum(1 for word in formal_indicators if word in full_text.lower())
        if formal_count >= 3:
            style_analysis['formality_level'] = 'formal'
        elif '!' in full_text or full_text.count('we') > full_text.count('shall'):
            style_analysis['formality_level'] = 'professional'

        # Calculate average word ranges for each section type
        for section_type, word_counts in style_analysis['section_word_ranges'].items():
            if word_counts:
                avg_words = sum(word_counts) / len(word_counts)
                min_words = min(word_counts)
                max_words = max(word_counts)
                style_analysis['section_word_ranges'][section_type] = {
                    'average': round(avg_words),
                    'range': [min_words, max_words],
                    'target': round(avg_words),
                    'flexibility': 0.2  # ±20% flexibility for matching template
                }

        print(f"📊 Style Analysis Complete:")
        print(f"  - Total words: {total_words}")
        print(f"  - Avg sentence length: {style_analysis['avg_sentence_length']} words")
        print(f"  - Formality: {style_analysis['formality_level']}")
        print(f"  - Sections with tables: {len(table_patterns)}")
        print(f"  - Writing samples collected: {len(writing_samples)}")

        return style_analysis

    def categorize_section(self, title: str) -> str:
        """Categorize section by title to understand typical word counts"""
        title_lower = title.lower()

        if any(word in title_lower for word in ['executive', 'summary', 'overview']):
            return 'executive_summary'
        elif any(word in title_lower for word in ['scope', 'approach', 'methodology']):
            return 'scope_approach'
        elif any(word in title_lower for word in ['technical', 'solution', 'architecture']):
            return 'technical_solution'
        elif any(word in title_lower for word in ['team', 'personnel', 'resources']):
            return 'team_resources'
        elif any(word in title_lower for word in ['timeline', 'schedule', 'milestones']):
            return 'timeline_schedule'
        elif any(word in title_lower for word in ['budget', 'cost', 'pricing', 'investment']):
            return 'budget_pricing'
        elif any(word in title_lower for word in ['compliance', 'requirements', 'standards']):
            return 'compliance_requirements'
        elif any(word in title_lower for word in ['experience', 'qualifications', 'references']):
            return 'experience_qualifications'
        else:
            return 'other'

    def create_ai_writing_guidelines(self, sections: List[Dict[str, Any]], style_analysis: Dict[str, Any]) -> Dict[str, Any]:
        """Create DETAILED guidelines for AI to EXACTLY mimic the template's writing style"""
        guidelines = {
            'structural_template': [],
            'writing_style_rules': [],
            'section_templates': {},
            'tone_guidelines': [],
            'critical_instructions': []  # Most important rules
        }

        # CRITICAL INSTRUCTIONS - These must be followed
        total_words = style_analysis.get('total_word_count', 1000)
        avg_sentence_length = style_analysis.get('avg_sentence_length', 15)
        formality = style_analysis.get('formality_level', 'professional')

        guidelines['critical_instructions'] = [
            f"MANDATORY: Match the {formality} tone exactly",
            f"MANDATORY: Target section word counts precisely (±20% max)",
            f"MANDATORY: Average sentence length must be {avg_sentence_length} words",
            f"MANDATORY: Use table/image placeholders where template has them",
            f"MANDATORY: Follow the exact section hierarchy and structure"
        ]

        # Create detailed structural template with subsections
        for section in sections:
            section_word_count = section.get('word_count', 200)
            section_has_tables = section.get('has_tables', False)
            section_sample = section.get('content_sample', '')

            section_template = {
                'title': section['title'],
                'level': section.get('level', 1),
                'target_word_count': section_word_count,
                'min_words': max(50, int(section_word_count * 0.8)),
                'max_words': int(section_word_count * 1.2),
                'has_tables': section_has_tables,
                'table_count': section.get('table_count', 0),
                'table_positions': section.get('tables', []),
                'writing_sample': section_sample[:300],  # First 300 chars as style reference
                'subsections': []
            }

            # Add subsection details
            for subsection in section.get('subsections', []):
                sub_word_count = subsection.get('word_count', 100)
                sub_has_tables = subsection.get('has_tables', False)
                sub_sample = subsection.get('content_sample', '')

                section_template['subsections'].append({
                    'title': subsection['title'],
                    'level': subsection.get('level', 2),
                    'target_word_count': sub_word_count,
                    'min_words': max(30, int(sub_word_count * 0.8)),
                    'max_words': int(sub_word_count * 1.2),
                    'has_tables': sub_has_tables,
                    'table_count': subsection.get('table_count', 0),
                    'writing_sample': sub_sample[:300]
                })

            guidelines['structural_template'].append(section_template)

            # Create section-specific template
            section_type = self.categorize_section(section['title'])
            if section_type not in guidelines['section_templates']:
                guidelines['section_templates'][section_type] = {
                    'word_count_target': section_word_count,
                    'typical_subsections': [sub['title'] for sub in section.get('subsections', [])],
                    'has_tables': section_has_tables,
                    'writing_example': section_sample[:500],
                    'style_instructions': f"Write like the '{section['title']}' section in the template"
                }

        # Writing style rules based on analysis
        guidelines['writing_style_rules'] = [
            f"Sentence length: Average {avg_sentence_length} words per sentence",
            f"Formality level: {formality.upper()} - maintain this throughout",
            f"Total document: Approximately {total_words} words",
            "Paragraph structure: Match the template's paragraph organization",
            "Technical depth: Match the level of technical detail in the template"
        ]

        # Tone guidelines with actual examples
        writing_samples = style_analysis.get('writing_samples', {})
        if writing_samples:
            sample_texts = list(writing_samples.values())[:3]
            guidelines['tone_guidelines'] = [
                f"Voice: {formality}, similar to these template examples:",
                *[f"Example {i+1}: {text[:150]}..." for i, text in enumerate(sample_texts) if text]
            ]

        # Add common phrases if found
        if style_analysis.get('common_phrases'):
            phrases = style_analysis['common_phrases']
            guidelines['tone_guidelines'].append(f"Use phrases like: {', '.join(phrases[:5])}")

        # Add table usage instructions
        table_patterns = style_analysis.get('table_usage_patterns', {})
        if table_patterns:
            guidelines['table_instructions'] = {
                'sections_with_tables': list(table_patterns.keys()),
                'usage': 'Insert table placeholders [TABLE: description] where template has tables',
                'patterns': table_patterns
            }

        print(f"📋 AI Writing Guidelines Created:")
        print(f"  - Structural template: {len(guidelines['structural_template'])} sections")
        print(f"  - Critical instructions: {len(guidelines['critical_instructions'])}")
        print(f"  - Section templates: {len(guidelines['section_templates'])} types")
        print(f"  - Writing samples: {len(writing_samples)}")

        return guidelines


def learn_toc_from_file(file_path: str, filename: str, template_name: str) -> Dict[str, Any]:
    """
    Main function to extract comprehensive template information from a DOCX file
    """
    try:
        if not DOCX_AVAILABLE:
            return {
                "error": "python-docx library not available. Cannot extract TOC from DOCX files."
            }

        if not os.path.exists(file_path):
            return {
                "error": f"File not found: {file_path}"
            }

        # Initialize analyzer
        analyzer = TemplateAnalyzer()

        # Load document
        print(f"📖 Loading document: {filename}")
        doc = Document(file_path)

        # Extract TOC structure
        print("🔍 Extracting TOC structure...")
        sections = analyzer.extract_toc_structure(doc)

        if not sections:
            return {
                "error": "No Word heading styles found. Please ensure your document uses Heading 1, Heading 2, or Heading 3 styles for section titles."
            }

        print(f"✅ Extracted {len(sections)} sections")

        # DETAILED DEBUG OUTPUT
        print("\n" + "="*80)
        print("COMPLETE EXTRACTED STRUCTURE DEBUG")
        print("="*80)

        for i, section in enumerate(sections):
            print(f"\nSECTION {i+1}:")
            print(f"  Title: '{section.get('title', 'NO TITLE')}'")
            print(f"  Level: {section.get('level', 'NO LEVEL')}")
            print(f"  Word Count: {section.get('word_count', 0)}")
            print(f"  Has Subsections: {bool(section.get('subsections'))}")

            if section.get('subsections'):
                print(f"  Subsections ({len(section['subsections'])}):")
                for j, subsection in enumerate(section['subsections']):
                    print(f"    {i+1}.{j+1}: '{subsection.get('title', 'NO TITLE')}' (Level {subsection.get('level', 'NO LEVEL')})")
            else:
                print(f"  Subsections: NONE")

        total_subsections = sum(len(section.get('subsections', [])) for section in sections)
        print(f"\nSUMMARY:")
        print(f"  Total Main Sections: {len(sections)}")
        print(f"  Total Subsections: {total_subsections}")
        print(f"  Total All Sections: {len(sections) + total_subsections}")
        print("="*80)

        # Analyze writing style
        print("🎨 Analyzing writing style...")
        style_analysis = analyzer.analyze_writing_style(sections)

        # Create AI writing guidelines
        print("🤖 Creating AI writing guidelines...")
        ai_guidelines = analyzer.create_ai_writing_guidelines(sections, style_analysis)

        # Create comprehensive template
        template = {
            "id": f"custom-{datetime.now().strftime('%Y%m%d-%H%M%S')}",
            "name": template_name,
            "template_name": template_name,
            "filename": filename,
            "category": "Custom",
            "description": f"Template extracted from {filename}",
            "created_at": datetime.now().isoformat(),

            # Core TOC structure
            "toc": [section['title'] for section in sections],
            "sections": [section['title'] for section in sections],
            "detailed_sections": sections,

            # Writing style and patterns
            "style_analysis": style_analysis,
            "ai_writing_guidelines": ai_guidelines,

            # Statistics for reference
            "statistics": {
                "total_sections": len(sections),
                "total_subsections": sum(len(section.get('subsections', [])) for section in sections),
                "total_words": style_analysis.get('total_word_count', 0),
                "avg_words_per_section": round(style_analysis.get('total_word_count', 0) / len(sections)) if sections else 0,
                "section_word_ranges": style_analysis.get('section_word_ranges', {}),
                "hierarchy_depth": max((section.get('level', 1) for section in sections), default=1)
            },

            # Preview for UI
            "preview": f"Custom template with {len(sections)} main sections, {sum(len(section.get('subsections', [])) for section in sections)} subsections, {style_analysis.get('total_word_count', 0)} words total"
        }

        # Save template to file
        template_file = analyzer.templates_dir / f"{template['id']}.json"
        with open(template_file, 'w', encoding='utf-8') as f:
            json.dump(template, f, indent=2, ensure_ascii=False)

        print(f"💾 Template saved: {template_file}")

        # Ingest template into ChromaDB with rich metadata
        print(f"📚 Ingesting template into ChromaDB with rich metadata...")
        _ingest_template_to_chromadb(template, file_path)

        return template

    except Exception as e:
        print(f"❌ Error in learn_toc_from_file: {e}")
        return {
            "error": f"Failed to extract TOC: {str(e)}"
        }


def get_saved_templates() -> List[Dict[str, Any]]:
    """Get all saved TOC templates"""
    templates_dir = Path("toc_templates")
    if not templates_dir.exists():
        return []

    templates = []
    for template_file in templates_dir.glob("*.json"):
        try:
            with open(template_file, 'r', encoding='utf-8') as f:
                template = json.load(f)
                templates.append(template)
        except Exception as e:
            print(f"⚠️ Error loading template {template_file}: {e}")

    return sorted(templates, key=lambda x: x.get('created_at', ''), reverse=True)


def get_template_by_id(template_id: str) -> Optional[Dict[str, Any]]:
    """Get a specific template by ID"""
    templates_dir = Path("toc_templates")
    template_file = templates_dir / f"{template_id}.json"

    if not template_file.exists():
        return None

    try:
        with open(template_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"⚠️ Error loading template {template_id}: {e}")
        return None


def delete_template_by_id(template_id: str) -> Dict[str, Any]:
    """Delete a template and all its associated data"""
    try:
        templates_dir = Path("toc_templates")
        template_file = templates_dir / f"{template_id}.json"

        if not template_file.exists():
            return {
                "status": "error",
                "message": f"Template {template_id} not found"
            }

        # Load template to get filename for cleanup
        template = get_template_by_id(template_id)

        # Delete the template file
        template_file.unlink()
        print(f"🗑️ Deleted template file: {template_file}")

        # Delete associated uploaded file if it exists
        if template and 'filename' in template:
            upload_file = Path("uploads") / template['filename']
            if upload_file.exists():
                upload_file.unlink()
                print(f"🗑️ Deleted upload file: {upload_file}")

        return {
            "status": "success",
            "message": f"Template {template_id} and associated data deleted successfully"
        }

    except Exception as e:
        print(f"❌ Error deleting template {template_id}: {e}")
        return {
            "status": "error",
            "message": f"Failed to delete template: {str(e)}"
        }


def _ingest_template_to_chromadb(template: Dict[str, Any], file_path: str):
    """
    Ingest template into ChromaDB with rich metadata for better retrieval.

    Metadata includes:
    - source: document filename
    - document_type: "template"
    - section: main section name
    - subsection: subsection name (if applicable)
    - level: hierarchy level (1 for main, 2 for subsection)
    - word_count: section word count
    - has_tables: whether section has tables
    - has_images: whether section has images
    - template_name: name of the template
    - template_id: unique template ID
    """
    try:
        from langchain_community.document_loaders import UnstructuredWordDocumentLoader
        from langchain.text_splitter import RecursiveCharacterTextSplitter
        from langchain.schema import Document as LangChainDocument
        from db import add_documents

        print(f"📚 Starting template ingestion for: {template['name']}")

        # Load the full document
        loader = UnstructuredWordDocumentLoader(file_path)
        full_docs = loader.load()

        if not full_docs:
            print("⚠️ No content loaded from document")
            return

        # Create documents with rich metadata for each section
        enriched_docs = []

        for section in template.get('detailed_sections', []):
            section_title = section.get('title', '')
            section_level = section.get('level', 1)
            section_word_count = section.get('word_count', 0)
            section_has_tables = section.get('has_tables', False)
            section_has_images = section.get('has_images', False)
            section_content = section.get('content_sample', '')

            # Create base metadata
            base_metadata = {
                "source": template['filename'],
                "document_type": "template",
                "template_name": template['name'],
                "template_id": template['id'],
                "section": section_title,
                "level": section_level,
                "word_count": section_word_count,
                "has_tables": section_has_tables,
                "has_images": section_has_images,
            }

            # Add main section document
            if section_content:
                enriched_docs.append(
                    LangChainDocument(
                        page_content=f"Section: {section_title}\n\n{section_content}",
                        metadata={**base_metadata, "subsection": None}
                    )
                )

            # Add subsection documents
            for subsection in section.get('subsections', []):
                subsection_title = subsection.get('title', '')
                subsection_level = subsection.get('level', 2)
                subsection_word_count = subsection.get('word_count', 0)
                subsection_has_tables = subsection.get('has_tables', False)
                subsection_has_images = subsection.get('has_images', False)
                subsection_content = subsection.get('content_sample', '')

                if subsection_content:
                    enriched_docs.append(
                        LangChainDocument(
                            page_content=f"Section: {section_title} > {subsection_title}\n\n{subsection_content}",
                            metadata={
                                **base_metadata,
                                "subsection": subsection_title,
                                "level": subsection_level,
                                "word_count": subsection_word_count,
                                "has_tables": subsection_has_tables,
                                "has_images": subsection_has_images,
                            }
                        )
                    )

        # Split large sections into chunks while preserving metadata
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

        chunked_docs = []
        for doc in enriched_docs:
            chunks = text_splitter.split_text(doc.page_content)
            for i, chunk in enumerate(chunks):
                chunked_docs.append(
                    LangChainDocument(
                        page_content=chunk,
                        metadata={
                            **doc.metadata,
                            "chunk_index": i,
                            "total_chunks": len(chunks)
                        }
                    )
                )

        # Ingest into "templates" collection
        if chunked_docs:
            result = add_documents(chunked_docs, collection="templates")
            print(f"✅ Ingested {len(chunked_docs)} template chunks into ChromaDB")
            print(f"   - Template: {template['name']}")
            print(f"   - Sections: {len(template.get('detailed_sections', []))}")
            print(f"   - Total chunks: {len(chunked_docs)}")
        else:
            print("⚠️ No content to ingest")

    except Exception as e:
        print(f"❌ Error ingesting template to ChromaDB: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    # Test the extractor
    test_file = "test_proposal.docx"
    if os.path.exists(test_file):
        result = learn_toc_from_file(test_file, test_file, "Test Template")
        print("\n📋 Extraction Result:")
        print(json.dumps(result, indent=2))
    else:
        print("No test file found. Place a test_proposal.docx file to test extraction.")