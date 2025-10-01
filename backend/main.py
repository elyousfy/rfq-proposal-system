# backend/main.py

import os
import json
import re
from typing import List

# Custom JSON encoder to handle bytes objects safely
class SafeJSONEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, bytes):
            try:
                return obj.decode('utf-8', errors='replace')
            except:
                return f"<binary data {len(obj)} bytes>"
        return super().default(obj)

from fastapi import FastAPI, UploadFile, File, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.exceptions import RequestValidationError
from pydantic import BaseModel
from starlette.exceptions import HTTPException as StarletteHTTPException
from dotenv import load_dotenv
from openai import OpenAI

from ingestion import ingest_paths
from retrieval import ask_question, format_context
from db import (
    search,
    safe_collection_name,
    delete_documents,
    get_rfqs,
    get_db_folders,
    get_chroma,
    load_data,
    save_data,
    add_rfq,
    add_file_to_rfq,
    drop_collection,
    inspect_collection,
)
from prompt import RFQ_EVALUATOR_PROMPT, RFQ_METADATA_PROMPT
from utils import file_to_text
from generation_control import controller, GenerationStatus

# --- Setup ---
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

app = FastAPI(
    title="RFQ / RFP QA Backend",
    description="Handles document ingestion, Q&A, and evaluation",
    version="0.3.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "./uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Global exception handlers for various errors
@app.exception_handler(UnicodeDecodeError)
async def unicode_decode_error_handler(request: Request, exc: UnicodeDecodeError):
    print(f"Unicode decode error: {exc}")
    return JSONResponse(
        status_code=400,
        content={
            "detail": "File contains invalid characters that cannot be processed. Please ensure the file is properly encoded in UTF-8 or use a different file format.",
            "error_type": "encoding_error"
        }
    )

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    print(f"Validation error: {exc}")

    # Check if this is likely a file upload Unicode issue
    error_str = str(exc).lower()
    if any(keyword in error_str for keyword in ["unicode", "decode", "utf-8", "invalid start byte"]):
        return JSONResponse(
            status_code=400,
            content={
                "detail": "File upload contains invalid characters. Please ensure files are properly encoded or use supported file formats (PDF, DOCX, TXT with UTF-8 encoding).",
                "error_type": "file_encoding_error"
            }
        )

    # For other validation errors, return a clean error without problematic bytes
    try:
        error_details = []
        for error in exc.errors():
            clean_error = {
                "type": error.get("type", "validation_error"),
                "loc": error.get("loc", []),
                "msg": str(error.get("msg", "Validation error"))
            }
            # Avoid including 'input' field which might contain problematic bytes
            error_details.append(clean_error)

        return JSONResponse(
            status_code=422,
            content={"detail": error_details}
        )
    except Exception as clean_error:
        print(f"Error cleaning validation error: {clean_error}")
        return JSONResponse(
            status_code=400,
            content={
                "detail": "Request validation failed. Please check your file uploads and try again.",
                "error_type": "validation_error"
            }
        )

@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    # Handle any remaining Unicode/encoding related errors
    if "unicode" in str(exc).lower() or "decode" in str(exc).lower() or "utf-8" in str(exc).lower():
        print(f"General encoding error: {exc}")
        return JSONResponse(
            status_code=400,
            content={
                "detail": "File processing error due to encoding issues. Please ensure files are properly formatted and encoded.",
                "error_type": "encoding_error"
            }
        )

    # For non-encoding errors, log and return generic error
    print(f"Unhandled exception: {exc}")
    return JSONResponse(
        status_code=500,
        content={
            "detail": "Internal server error. Please try again.",
            "error_type": "server_error"
        }
    )

# --- Models ---

class AskRequest(BaseModel):
    question: str
    top_k: int = 6

class EvaluateRFQRequest(BaseModel):
    rfqName: str
    client: str
    dueDate: str
    documents: List[str]

class DeleteRequest(BaseModel):
    collection: str
    filename: str

class SaveRFQRequest(BaseModel):
    name: str
    client: str
    dueDate: str
    mainDocument: str
    supportingDocuments: list[str] = []

class DeleteRFQRequest(BaseModel):
    name: str
    
class AddSupportingDocRequest(BaseModel):
    rfqName: str
    filename: str

class ReplaceMainDocRequest(BaseModel):
    rfqName: str
    oldFilename: str
    newFilename: str

class CreateFolderRequest(BaseModel):
    name: str

class SaveEvaluationRequest(BaseModel):
    rfqName: str
    evaluation: dict
    valueUSD: float

class GenerateProposalRequest(BaseModel):
    rfqName: str
    requirements: list[str] = []
    structure: str = "standard"  # standard, technical, services
    tone: str = "professional"  # professional, formal, innovative
    includeCompliance: bool = True
    tocTemplateId: str = None  # Optional TOC template ID
    sessionId: str = None  # Optional session ID for pause/stop/resume control

class GenerateSectionRequest(BaseModel):
    rfqName: str
    sectionType: str  # executive_summary, technical_approach, etc.
    context: str = ""
    requirements: list[str] = []
    tone: str = "professional"

class ComplianceMatrixRequest(BaseModel):
    rfqName: str
    requirements: list[str] = []

class ProposalTemplate(BaseModel):
    name: str
    industry: str
    sections: list[dict]
    variables: list[dict]

class LearnProposalRequest(BaseModel):
    filename: str
    client_type: str = ""
    proposal_type: str = ""

class GenerateFromTemplateRequest(BaseModel):
    rfqName: str
    client_type: str
    proposal_type: str = ""
    requirements: list[str] = []

class ExtractTOCRequest(BaseModel):
    filename: str
    template_name: str = ""

class ApplyTOCTemplateRequest(BaseModel):
    template_id: str
    proposal_title: str = "New Proposal"


# --- Routes ---

@app.get("/")
async def root():
    return {
        "status": "ok",
        "endpoints": [
            "/ingest", "/ask", "/evaluate_rfq", "/rfqs", "/database",
            "/extract_rfq_metadata", "/save_rfq", "/replace_main_doc",
            "/add_supporting_doc", "/delete", "/delete_rfq", "/download/{filename}",
            "/inspect/{collection}", "/test_embeddings", "/create_folder",
            "/save_evaluation", "/get_evaluation/{rfq_name}", "/get_all_evaluations",
            "/generate_proposal", "/generate_section", "/generate_compliance_matrix",
            "/get_proposal_templates", "/export_proposal/{format}",
            "/learn_proposal", "/generate_from_template", "/get_learned_templates",
            "/extract_toc", "/get_toc_templates", "/apply_toc_template", "/get_toc_preview/{template_id}", "/delete_template/{template_id}",
        ]
    }

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

@app.post("/ingest")
async def ingest(files: List[UploadFile] = File(...), collection: str = "global"):
    from db import add_file_to_folder
    from datetime import datetime

    collection = safe_collection_name(collection)
    saved_paths = []
    for file in files:
        try:
            # Read file content safely
            file_content = await file.read()

            save_path = os.path.join(UPLOAD_DIR, file.filename)
            with open(save_path, "wb") as f:
                f.write(file_content)
            saved_paths.append(save_path)
        except Exception as file_error:
            print(f"Error processing file {file.filename}: {file_error}")
            # Continue with other files instead of failing completely
            continue

    n_chunks = ingest_paths(saved_paths, upload_dir=UPLOAD_DIR, collection=collection)
    
    # If this is a database folder upload, track files in metadata
    if collection.startswith("db_"):
        # Find the actual folder name by comparing with existing database folders
        data = load_data()
        target_folder = None
        
        for folder in data.get("database", []):
            if safe_collection_name(f"db_{folder['name']}") == collection:
                target_folder = folder["name"]
                break
        
        if target_folder:
            for file in files:
                file_entry = {
                    "name": file.filename,
                    "uploadedAt": datetime.now().isoformat()
                }
                add_file_to_folder(target_folder, file_entry)
    
    return {"status": "success", "collection": collection, "files": saved_paths, **n_chunks}

@app.post("/ask")
async def ask(request: AskRequest):
    return ask_question(request.question, k=request.top_k)

@app.post("/evaluate_rfq")
async def evaluate_rfq(request: EvaluateRFQRequest):
    print(f"🎯 Evaluating RFQ: {request.rfqName}")
    print(f"📊 Request data: {request}")
    
    collection_name = safe_collection_name(f"rfq_{request.rfqName}")
    print(f"🗂️ Using collection: {collection_name}")
    
    query = f"Extract key requirements and context from RFQ {request.rfqName}"
    docs = search(query, k=15, collection=collection_name)
    print(f"📄 Retrieved {len(docs)} documents from collection")
    
    if not docs:
        print("⚠️ No documents found in collection!")
        return {k: [] for k in ["objectives", "deliverables", "constraints", "risks",
                                "successCriteria", "stakeholders", "standards", "scope"]}
    
    context = format_context(docs)
    print(f"📝 Context length: {len(context)} characters")
    print(f"📋 First 500 chars of context: {context[:500]}...")

    messages = [
        {"role": "system", "content": RFQ_EVALUATOR_PROMPT},
        {"role": "user", "content": f"RFQ Context:\n{context}"},
    ]

    print("🤖 Sending request to OpenAI...")
    response = client.chat.completions.create(
        model=os.getenv("OPENAI_MODEL", "gpt-4o"),
        messages=messages,
        temperature=0,
    )

    raw = response.choices[0].message.content.strip()
    print(f"📤 OpenAI response: {raw}")
    
    # Clean up markdown code blocks and extra text
    import re
    print(f"🔍 Raw response first 100 chars: {repr(raw[:100])}")
    print(f"🔍 Raw response last 100 chars: {repr(raw[-100:])}")
    
    # Try multiple cleaning strategies
    clean_json = raw.strip()
    
    # Remove markdown code blocks (comprehensive approach)
    if '```' in clean_json:
        print("🧹 Removing markdown code blocks...")
        # Find JSON content between code blocks
        json_match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', clean_json, re.DOTALL)
        if json_match:
            clean_json = json_match.group(1).strip()
        else:
            # Fallback: remove all lines with ```
            lines = clean_json.split('\n')
            lines = [line for line in lines if not line.strip().startswith('```')]
            clean_json = '\n'.join(lines)
    
    # Remove any leading/trailing non-JSON text
    # Look for the first { and last }
    start_idx = clean_json.find('{')
    end_idx = clean_json.rfind('}')
    
    if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
        clean_json = clean_json[start_idx:end_idx+1]
    
    clean_json = clean_json.strip()
    print(f"🧹 Cleaned JSON first 200 chars: {clean_json[:200]}...")
    print(f"🧹 Cleaned JSON last 100 chars: {clean_json[-100:]}")
    
    try:
        parsed = json.loads(clean_json)
        print("✅ Successfully parsed JSON response")
        print(f"📊 Parsed data keys: {list(parsed.keys())}")
        return parsed
    except Exception as e:
        print(f"❌ Failed to parse JSON: {e}")
        print(f"❌ Problematic JSON content: {repr(clean_json[:500])}")
        return {k: [] for k in ["objectives", "deliverables", "constraints", "risks",
                                "successCriteria", "stakeholders", "standards", "scope"]}

@app.post("/extract_rfq_metadata")
async def extract_rfq_metadata(file: UploadFile = File(...)):
    try:
        contents = await file.read()

        # Pre-validate that we can handle the file content
        try:
            # Try to detect encoding and handle invalid bytes early
            if file.filename and file.filename.lower().endswith('.txt'):
                import chardet
                detected = chardet.detect(contents)
                encoding = detected.get("encoding", "utf-8")
                # Test decode with error handling
                contents.decode(encoding, errors="replace")
        except Exception as decode_error:
            print(f"File encoding validation warning for {file.filename}: {decode_error}")

        text = file_to_text(contents, file.filename)

        if not text.strip():
            return {"status": "error", "metadata": {"name": "", "client": "", "dueDate": ""}}

        snippet = text[:10000] + "\n...\n" + text[-10000:]
        messages = [
            {"role": "system", "content": RFQ_METADATA_PROMPT},
            {"role": "user", "content": f"Extract metadata from this RFQ:\n{snippet}"},
        ]

        response = client.chat.completions.create(
            model=os.getenv("OPENAI_MODEL", "gpt-4o"),
            messages=messages,
            temperature=0,
        )

        raw = response.choices[0].message.content.strip()
        raw_clean = re.sub(r"^```(?:json)?|```$", "", raw, flags=re.MULTILINE).strip()

        try:
            data = json.loads(raw_clean)
        except Exception:
            data = {"name": "", "client": "", "dueDate": ""}

        return {
            "status": "success",
            "metadata": {
                "name": data.get("name", ""),
                "client": data.get("client", ""),
                "dueDate": data.get("dueDate", ""),
            },
        }
    except Exception as e:
        print(f"Metadata extraction error: {e}")
        return {"status": "error", "metadata": {"name": "", "client": "", "dueDate": ""}}

@app.post("/save_rfq")
async def save_rfq(request: SaveRFQRequest):
    rfq_entry = {
        "name": request.name,
        "client": request.client,
        "dueDate": request.dueDate,
        "mainDocument": request.mainDocument,
        "supportingDocuments": request.supportingDocuments,
        "documents": [request.mainDocument] + request.supportingDocuments,
    }
    saved = add_rfq(rfq_entry)

    collection = safe_collection_name(f"rfq_{request.name}")
    paths = [os.path.join(UPLOAD_DIR, f) for f in rfq_entry["documents"] if f]
    if paths:
        ingest_paths(paths, upload_dir=UPLOAD_DIR, collection=collection)

    return {"status": "success", "rfq": saved}

@app.post("/replace_main_doc")
async def replace_main_doc(req: ReplaceMainDocRequest):
    collection = safe_collection_name(f"rfq_{req.rfqName}")
    delete_documents(collection, req.oldFilename)

    old_path = os.path.join(UPLOAD_DIR, req.oldFilename)
    if os.path.exists(old_path):
        os.remove(old_path)

    new_path = os.path.join(UPLOAD_DIR, req.newFilename)
    ingest_paths([new_path], upload_dir=UPLOAD_DIR, collection=collection)

    data = load_data()
    for rfq in data["rfqs"]:
        if rfq["name"] == req.rfqName:
            rfq["mainDocument"] = req.newFilename
            rfq["documents"] = [req.newFilename] + rfq.get("supportingDocuments", [])
    save_data(data)

    return {"status": "success", "rfq": req.rfqName, "newMain": req.newFilename}

@app.post("/add_supporting_doc")
async def add_supporting_doc(req: AddSupportingDocRequest):
    """
    Register a supporting document into the RFQ's metadata.
    Note: Ingestion should be done separately via /ingest endpoint.
    """
    result = add_file_to_rfq(req.rfqName, req.filename)
    if result.get("status") == "error":
        return {"status": "error", "message": result["reason"]}

    print(f"📝 Registered {req.filename} as supporting doc for RFQ: {req.rfqName}")
    return {"status": "success", "rfq": req.rfqName, "file": req.filename}
@app.post("/delete")
async def delete_file(req: DeleteRequest):
    collection = safe_collection_name(req.collection)
    delete_documents(collection, req.filename)

    file_path = os.path.join(UPLOAD_DIR, req.filename)
    if os.path.exists(file_path):
        os.remove(file_path)

    data = load_data()
    for rfq in data.get("rfqs", []):
        if rfq.get("mainDocument") == req.filename:
            rfq["mainDocument"] = None
        if req.filename in rfq.get("supportingDocuments", []):
            rfq["supportingDocuments"].remove(req.filename)
    save_data(data)

    return {"status": "success", "deleted": req.filename, "collection": collection}

@app.post("/delete_rfq")
async def delete_rfq(req: DeleteRFQRequest):
    """
    Delete an entire RFQ:
    - remove metadata
    - delete all uploaded files
    - drop its Chroma collection
    """
    data = load_data()

    # find the RFQ entry to know which files to delete
    rfq_entry = next((r for r in data.get("rfqs", []) if r["name"] == req.name), None)
    if rfq_entry:
        all_docs = [rfq_entry.get("mainDocument")] + rfq_entry.get("supportingDocuments", [])
        for doc in filter(None, all_docs):
            file_path = os.path.join(UPLOAD_DIR, doc)
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"⚠️ Could not delete file {file_path}: {e}")

    # remove from metadata
    data["rfqs"] = [r for r in data.get("rfqs", []) if r["name"] != req.name]
    save_data(data)

    # drop the whole collection in Chroma
    collection = safe_collection_name(f"rfq_{req.name}")
    dropped = drop_collection(collection)
    if not dropped:
        print(f"⚠️ Failed to fully drop collection: {collection}")

    return {"status": "success", "deleted": req.name, "collection": collection, "collectionDropped": dropped}

@app.get("/rfqs")
async def list_rfqs():
    return get_rfqs()

@app.get("/database")
async def list_database():
    return get_db_folders()

@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(file_path, filename=filename)

@app.get("/view/{filename}")
async def view_file(filename: str):
    """
    View a file inline (for iframe viewing) without triggering download.
    """
    file_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found")
    
    # Determine content type based on file extension
    import mimetypes
    content_type, _ = mimetypes.guess_type(file_path)
    
    # For PDFs, set inline disposition
    if filename.lower().endswith('.pdf'):
        content_type = "application/pdf"
    
    return FileResponse(
        file_path, 
        media_type=content_type,
        headers={"Content-Disposition": "inline"}
    )

@app.get("/inspect/{collection}")
async def inspect_collection_endpoint(collection: str):
    """
    Inspect a collection to verify chunks, embeddings, and ingestion quality.
    Example: GET /inspect/rfq_Client_A___Security_RFP
    """
    result = inspect_collection(collection)
    return result

@app.get("/test_embeddings")
async def test_embeddings():
    """
    Test OpenAI embeddings connection and functionality.
    """
    try:
        from db import embeddings
        import time
        
        test_text = "This is a test document to verify embeddings are working correctly."
        
        print(f"🧠 Testing OpenAI embeddings connection...")
        start_time = time.time()
        
        # Test single embedding
        embedding = embeddings.embed_query(test_text)
        
        end_time = time.time()
        duration = end_time - start_time
        
        # Test batch embedding
        test_batch = [test_text, "Another test text for batch processing"]
        batch_embeddings = embeddings.embed_documents(test_batch)
        
        import numpy as np
        embedding_norm = np.linalg.norm(embedding)
        
        result = {
            "status": "success",
            "embedding_model": "text-embedding-3-small",
            "dimensions": len(embedding),
            "single_embedding": {
                "duration_seconds": round(duration, 3),
                "norm": round(float(embedding_norm), 3),
                "sample_values": [round(float(x), 6) for x in embedding[:5]]
            },
            "batch_embedding": {
                "batch_size": len(batch_embeddings),
                "all_same_dimension": all(len(emb) == len(embedding) for emb in batch_embeddings)
            },
            "test_text": test_text
        }
        
        print(f"✅ Embeddings test successful: {len(embedding)} dims, norm: {embedding_norm:.3f}")
        return result
        
    except Exception as e:
        error_result = {
            "status": "error",
            "message": str(e),
            "error_type": type(e).__name__
        }
        print(f"❌ Embeddings test failed: {e}")
        return error_result

@app.post("/create_folder")
async def create_folder(request: CreateFolderRequest):
    """
    Create a new database folder (just adds it to metadata for now).
    The actual vector collection will be created when files are uploaded.
    """
    data = load_data()
    
    # Check if folder already exists
    folder_names = [f.get("name", "") for f in data.get("database", [])]
    if request.name in folder_names:
        return {"status": "error", "message": "Folder already exists"}
    
    # Add new empty folder
    if "database" not in data:
        data["database"] = []
    
    data["database"].append({
        "name": request.name,
        "files": []
    })
    
    save_data(data)
    
    return {"status": "success", "folder": request.name}

@app.post("/save_evaluation")
async def save_evaluation(request: SaveEvaluationRequest):
    """
    Save an RFQ evaluation. Overwrites any existing evaluation for the same RFQ.
    """
    from datetime import datetime
    
    data = load_data()
    
    # Initialize evaluations section if it doesn't exist
    if "evaluations" not in data:
        data["evaluations"] = {}
    
    # Save the evaluation with timestamp
    data["evaluations"][request.rfqName] = {
        "evaluation": request.evaluation,
        "valueUSD": request.valueUSD,
        "timestamp": datetime.now().isoformat(),
        "rfqName": request.rfqName
    }
    
    save_data(data)
    
    return {"status": "success", "rfqName": request.rfqName}

@app.get("/get_evaluation/{rfq_name}")
async def get_evaluation(rfq_name: str):
    """
    Get saved evaluation for a specific RFQ.
    """
    data = load_data()
    evaluations = data.get("evaluations", {})
    
    if rfq_name in evaluations:
        return {
            "status": "success", 
            "evaluation": evaluations[rfq_name]
        }
    else:
        return {
            "status": "not_found",
            "message": f"No saved evaluation found for RFQ: {rfq_name}"
        }

@app.get("/get_all_evaluations")
async def get_all_evaluations():
    """
    Get all saved evaluations.
    """
    data = load_data()
    evaluations = data.get("evaluations", {})
    
    return {
        "status": "success",
        "evaluations": evaluations
    }

@app.post("/generate_proposal")
async def generate_proposal(request: GenerateProposalRequest):
    """
    Generate a sophisticated proposal using advanced generation system with current database.
    """
    from advanced_generator import generate_advanced_proposal
    from toc_extractor import get_toc_templates
    from datetime import datetime

    try:
        print(f"🤖 Generating advanced proposal for RFQ: {request.rfqName}")
        print(f"📋 Structure: {request.structure}, Tone: {request.tone}")

        # Get TOC template if provided
        toc_template = None
        if request.tocTemplateId:
            print(f"🎯 Using TOC Template: {request.tocTemplateId}")
            templates = get_toc_templates()
            toc_template = next((t for t in templates if t.get("id") == request.tocTemplateId), None)
            if not toc_template:
                print(f"⚠️ TOC Template {request.tocTemplateId} not found, using default structure")

        # Generate proposal using advanced system with current database
        proposal = generate_advanced_proposal(
            rfq_name=request.rfqName,
            toc_template=toc_template,
            tone=request.tone,
            top_k=6,
            session_id=request.sessionId  # Pass session_id for pause/stop/resume control
        )

        proposal["generated_at"] = datetime.now().isoformat()

        print(f"✅ Generated advanced proposal with {len(proposal['sections'])} sections")
        return {
            "status": "success",
            "proposal": proposal
        }

    except Exception as e:
        print(f"❌ Error generating advanced proposal: {e}")
        import traceback
        traceback.print_exc()
        return {
            "status": "error",
            "message": str(e)
        }

@app.post("/generate_section")
async def generate_section(request: GenerateSectionRequest):
    """
    Generate a specific proposal section using advanced generation system.
    """
    from advanced_generator import generate_advanced_section
    from db import safe_collection_name

    try:
        print(f"🤖 Generating advanced section: {request.sectionType} for RFQ: {request.rfqName}")

        collection_name = safe_collection_name(f"rfq_{request.rfqName}")

        section_result = generate_advanced_section(
            section_title=request.sectionType,
            rfq_collection=collection_name,
            level=1,
            outline_path=request.sectionType,
            top_k=6
        )

        print(f"✅ Generated advanced section: {request.sectionType}")
        return {
            "status": "success",
            "section": {
                "type": request.sectionType,
                "title": section_result.get("title", request.sectionType),
                "content": section_result.get("content", ""),
                "notes": section_result.get("notes", []),
                "risks": section_result.get("risks", []),
                "assumptions": section_result.get("assumptions", []),
                "image_suggestions": section_result.get("image_suggestions", [])
            }
        }

    except Exception as e:
        print(f"❌ Error generating advanced section: {e}")
        import traceback
        traceback.print_exc()
        return {
            "status": "error",
            "message": str(e)
        }

@app.get("/generation_status/{session_id}")
async def get_generation_status(session_id: str):
    """Get current generation status."""
    status = controller.get_status(session_id)
    if not status:
        raise HTTPException(status_code=404, detail="Session not found")
    return {
        "status": status['status'].value,
        "current_section": status['current_section'],
        "completed_sections": status['completed_sections'],
        "total_sections": status['total_sections'],
        "message": status['message']
    }

@app.post("/generation_pause/{session_id}")
async def pause_generation(session_id: str):
    """Pause proposal generation."""
    success = controller.pause(session_id)
    return {"status": "success" if success else "error", "paused": success}

@app.post("/generation_resume/{session_id}")
async def resume_generation(session_id: str):
    """Resume paused generation."""
    success = controller.resume(session_id)
    return {"status": "success" if success else "error", "resumed": success}

@app.post("/generation_stop/{session_id}")
async def stop_generation(session_id: str):
    """Stop proposal generation completely."""
    success = controller.stop(session_id)
    return {"status": "success" if success else "error", "stopped": success}

@app.post("/generate_compliance_matrix")
async def generate_compliance_matrix(request: ComplianceMatrixRequest):
    """
    Generate a compliance matrix for the RFQ requirements.
    """
    from proposal_generator import generate_compliance_matrix
    
    try:
        print(f"🤖 Generating compliance matrix for RFQ: {request.rfqName}")
        
        matrix = generate_compliance_matrix(
            rfq_name=request.rfqName,
            requirements=request.requirements
        )
        
        print(f"✅ Generated compliance matrix with {len(matrix)} items")
        return {
            "status": "success",
            "compliance_matrix": matrix
        }
        
    except Exception as e:
        print(f"❌ Error generating compliance matrix: {e}")
        return {
            "status": "error",
            "message": str(e)
        }

@app.get("/get_proposal_templates")
async def get_proposal_templates():
    """
    Get available proposal templates.
    """
    from proposal_generator import PROPOSAL_TEMPLATES
    
    return {
        "status": "success",
        "templates": PROPOSAL_TEMPLATES
    }


@app.post("/export_proposal/{format}")
async def export_proposal(format: str, proposal_data: dict):
    """
    Export proposal to PDF or DOCX format.
    """
    try:
        # Create formatted content with proper structure
        title = proposal_data.get('title', 'Untitled Proposal')
        rfq_name = proposal_data.get('rfqName', 'N/A')
        generated_date = proposal_data.get('updatedAt', 'N/A')
        sections = proposal_data.get('sections', [])
        
        if format.lower() == "pdf":
            # Create HTML content for PDF conversion
            html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Arial', sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
        }}
        h1 {{
            color: #2c3e50;
            border-bottom: 3px solid #3498db;
            padding-bottom: 10px;
            margin-bottom: 30px;
        }}
        h2 {{
            color: #34495e;
            border-bottom: 1px solid #bdc3c7;
            padding-bottom: 5px;
            margin-top: 30px;
            margin-bottom: 15px;
        }}
        .header-info {{
            background: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            margin-bottom: 30px;
        }}
        .header-info p {{
            margin: 5px 0;
        }}
        .section {{
            margin-bottom: 25px;
        }}
        ul, ol {{
            padding-left: 25px;
        }}
        blockquote {{
            border-left: 4px solid #3498db;
            margin: 20px 0;
            padding-left: 20px;
            font-style: italic;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 15px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
        }}
        @media print {{
            body {{ font-size: 12pt; }}
            h1 {{ font-size: 18pt; }}
            h2 {{ font-size: 14pt; }}
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    
    <div class="header-info">
        <p><strong>RFQ:</strong> {rfq_name}</p>
        <p><strong>Generated:</strong> {generated_date}</p>
        <p><strong>Document Type:</strong> Business Proposal</p>
    </div>
"""
            
            # Add sections with proper content field
            for section in sections:
                section_title = section.get('title', 'Untitled Section')
                section_content = section.get('content', '')

                html_content += f'\n    <div class="section">\n        <h2>{section_title}</h2>\n'

                # Convert content to HTML
                if section_content:
                    # Basic markdown to HTML conversion
                    lines = section_content.split('\n')
                    in_list = False
                    in_blockquote = False
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        if in_list:
                            html_content += '        </ul>\n'
                            in_list = False
                        if in_blockquote:
                            html_content += '        </blockquote>\n'
                            in_blockquote = False
                        html_content += '        <br/>\n'
                        continue
                    
                    # Handle headers
                    if line.startswith('### '):
                        html_content += f'        <h3>{line[4:]}</h3>\n'
                    elif line.startswith('## '):
                        html_content += f'        <h3>{line[3:]}</h3>\n'
                    elif line.startswith('# '):
                        html_content += f'        <h3>{line[2:]}</h3>\n'
                    # Handle lists
                    elif line.startswith('- ') or line.startswith('* '):
                        if not in_list:
                            html_content += '        <ul>\n'
                            in_list = True
                        html_content += f'            <li>{line[2:]}</li>\n'
                    # Handle blockquotes
                    elif line.startswith('> '):
                        if not in_blockquote:
                            html_content += '        <blockquote>\n'
                            in_blockquote = True
                        html_content += f'        <p>{line[2:]}</p>\n'
                    # Handle regular paragraphs
                    else:
                        if in_list:
                            html_content += '        </ul>\n'
                            in_list = False
                        if in_blockquote:
                            html_content += '        </blockquote>\n'
                            in_blockquote = False
                        
                        # Basic text formatting
                        formatted_line = line
                        formatted_line = formatted_line.replace('**', '<strong>').replace('**', '</strong>')
                        formatted_line = formatted_line.replace('*', '<em>').replace('*', '</em>')
                        
                        html_content += f'        <p>{formatted_line}</p>\n'
                
                # Close any open tags
                if in_list:
                    html_content += '        </ul>\n'
                if in_blockquote:
                    html_content += '        </blockquote>\n'
                    
                html_content += '    </div>\n'
            
            html_content += """
</body>
</html>"""
            
            return {
                "status": "success",
                "format": "html",
                "content": html_content,
                "filename": f"{title.replace(' ', '_')}.html"
            }
            
        elif format.lower() == "docx":
            # Generate proper DOCX file using python-docx
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            import io
            import base64
            import os
            import tempfile

            # Create a new Document
            doc = Document()

            # Set up styles
            styles = doc.styles

            # Modify existing styles
            title_style = styles['Heading 1']
            title_style.font.size = Pt(18)
            title_style.font.bold = True

            header_style = styles['Heading 2']
            header_style.font.size = Pt(14)
            header_style.font.bold = True

            subheader_style = styles['Heading 3']
            subheader_style.font.size = Pt(12)
            subheader_style.font.bold = True

            # Add document title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add document info
            info_para = doc.add_paragraph()
            info_para.add_run(f"RFQ: {rfq_name}").bold = True
            info_para.add_run(f"\nGenerated: {generated_date}")
            info_para.add_run(f"\nDocument Type: Business Proposal")

            # Add separator
            doc.add_paragraph("_" * 80)

            # Process sections with proper formatting
            for section in sections:
                section_title = section.get('title', 'Untitled Section')
                section_content = section.get('content', '')
                section_level = section.get('level', 1)

                # Add section header with appropriate level
                if section_level == 1:
                    doc.add_heading(section_title, level=1)
                elif section_level == 2:
                    doc.add_heading(section_title, level=2)
                else:
                    doc.add_heading(section_title, level=3)

                # Process section content with markdown parsing
                if section_content:
                    lines = section_content.split('\n')
                    current_para = None

                    for line in lines:
                        line = line.strip()

                        if line.startswith('### '):
                            # Subsubheading
                            doc.add_heading(line[4:], level=4)
                            current_para = None
                        elif line.startswith('## '):
                            # Subheading
                            doc.add_heading(line[3:], level=3)
                            current_para = None
                        elif line.startswith('# '):
                            # Heading
                            doc.add_heading(line[2:], level=2)
                            current_para = None
                        elif line.startswith('- ') or line.startswith('* '):
                            # Bullet point
                            bullet_para = doc.add_paragraph(line[2:], style='List Bullet')
                            current_para = None
                        elif line.startswith('> '):
                            # Quote
                            quote_para = doc.add_paragraph()
                            quote_run = quote_para.add_run(line[2:])
                            quote_run.italic = True
                            current_para = None
                        elif line.startswith('[TABLE:') or line.startswith('[IMAGE:'):
                            # Placeholder for tables/images
                            placeholder_para = doc.add_paragraph()
                            placeholder_run = placeholder_para.add_run(line)
                            placeholder_run.italic = True
                            placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            current_para = None
                        elif line:
                            # Regular text - handle bold/italic
                            if current_para is None:
                                current_para = doc.add_paragraph()

                            # Simple markdown processing for bold text
                            text = line
                            while '**' in text:
                                before, after = text.split('**', 1)
                                current_para.add_run(before)
                                if '**' in after:
                                    bold_text, text = after.split('**', 1)
                                    current_para.add_run(bold_text).bold = True
                                else:
                                    current_para.add_run('**' + after)
                                    text = ''
                            if text:
                                current_para.add_run(text)
                        else:
                            # Empty line - start new paragraph
                            current_para = None

                # Add spacing between sections
                doc.add_paragraph()

            # Save to temporary file and return base64 encoded content
            # Use BytesIO instead of temp file to avoid Windows file locking issues
            from io import BytesIO
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            docx_bytes = docx_buffer.read()
            docx_buffer.close()

            # Encode as base64 for transfer
            docx_base64 = base64.b64encode(docx_bytes).decode('utf-8')

            return {
                "status": "success",
                "format": "docx",
                "content": docx_base64,
                "filename": f"{title.replace(' ', '_').replace('/', '_')}.docx"
            }
        else:
            return {
                "status": "error",
                "message": "Unsupported format. Use 'pdf' or 'docx'"
            }
            
    except Exception as e:
        print(f"❌ Error exporting proposal: {e}")
        return {
            "status": "error",
            "message": str(e)
        }

@app.post("/learn_proposal")
async def learn_proposal(request: LearnProposalRequest):
    """
    Learn from an existing proposal to extract structure, tone, and patterns.
    """
    from proposal_learner import analyze_and_learn_proposal
    from datetime import datetime
    
    try:
        print(f"📚 Learning from proposal: {request.filename}")
        
        # Find the uploaded file
        file_path = os.path.join(UPLOAD_DIR, request.filename)
        if not os.path.exists(file_path):
            return {
                "status": "error",
                "message": f"File not found: {request.filename}"
            }
        
        # Analyze and learn from the proposal
        template = analyze_and_learn_proposal(
            file_path=file_path,
            filename=request.filename,
            client_type=request.client_type
        )
        
        if not template:
            return {
                "status": "error",
                "message": "Failed to analyze proposal. Please check file format and content."
            }
        
        # Add timestamp
        template["learned_at"] = datetime.now().isoformat()
        
        print(f"✅ Successfully learned from {request.filename}")
        return {
            "status": "success",
            "template": template,
            "message": f"Successfully learned proposal structure and style from {request.filename}"
        }
        
    except Exception as e:
        print(f"❌ Error learning from proposal: {e}")
        return {
            "status": "error",
            "message": str(e)
        }

@app.post("/generate_from_template")
async def generate_from_template(request: GenerateFromTemplateRequest):
    """
    Generate a new proposal based on learned templates.
    """
    from proposal_learner import generate_proposal_from_template
    from datetime import datetime
    
    try:
        print(f"🎯 Generating proposal from template for {request.client_type}")
        
        proposal = generate_proposal_from_template(
            rfq_name=request.rfqName,
            client_type=request.client_type,
            proposal_type=request.proposal_type,
            requirements=request.requirements
        )
        
        if "error" in proposal:
            return {
                "status": "error",
                "message": proposal["error"]
            }
        
        # Add timestamp
        proposal["generated_at"] = datetime.now().isoformat()
        
        print(f"✅ Generated proposal using learned templates")
        return {
            "status": "success",
            "proposal": proposal
        }
        
    except Exception as e:
        print(f"❌ Error generating from template: {e}")
        return {
            "status": "error",
            "message": str(e)
        }

@app.get("/get_learned_templates")
async def get_learned_templates():
    """
    Get all learned proposal templates.
    """
    try:
        data = load_data()
        templates = data.get("proposal_templates", [])
        
        # Add summary information
        summary = {
            "total_templates": len(templates),
            "client_types": list(set([t.get("client_type", "Unknown") for t in templates if t.get("client_type")])),
            "industries": list(set([t.get("industry", "Unknown") for t in templates if t.get("industry")])),
            "templates": templates
        }
        
        return {
            "status": "success",
            "data": summary
        }
        
    except Exception as e:
        print(f"❌ Error getting learned templates: {e}")
        return {
            "status": "error",
            "message": str(e),
            "data": {
                "total_templates": 0,
                "client_types": [],
                "industries": [],
                "templates": []
            }
        }

@app.post("/extract_toc")
async def extract_toc(request: ExtractTOCRequest):
    """
    Extract table of contents from an uploaded DOCX file and save as template.
    """
    from toc_extractor import learn_toc_from_file

    try:
        print(f"📑 Extracting TOC from: {request.filename}")

        # Find the uploaded file
        file_path = os.path.join(UPLOAD_DIR, request.filename)
        if not os.path.exists(file_path):
            return {
                "status": "error",
                "message": f"File not found: {request.filename}"
            }

        # Extract TOC and save as template
        result = learn_toc_from_file(file_path, request.filename, request.template_name)

        if "error" in result:
            return {
                "status": "error",
                "message": result["error"]
            }

        print(f"✅ TOC template created: {result['template_name']}")
        return {
            "status": "success",
            "template": result,
            "message": f"TOC template '{result['template_name']}' created successfully"
        }

    except Exception as e:
        print(f"❌ Error extracting TOC: {e}")
        return {
            "status": "error",
            "message": str(e)
        }

@app.get("/get_toc_templates")
async def get_toc_templates():
    """
    Get all available TOC templates.
    """
    from toc_extractor import get_saved_templates

    try:
        # Get custom templates from toc_extractor
        custom_templates = get_saved_templates()
        print(f"📋 Loaded {len(custom_templates)} custom templates")

        # DEBUG: Check if detailed_sections exists in first template
        if custom_templates:
            print(f"🔍 DEBUG First template keys: {list(custom_templates[0].keys())}")
            print(f"🔍 DEBUG Has detailed_sections: {('detailed_sections' in custom_templates[0])}")
            if 'detailed_sections' in custom_templates[0]:
                print(f"🔍 DEBUG detailed_sections length: {len(custom_templates[0]['detailed_sections'])}")

        # Predefined templates (fallback and examples)
        predefined_templates = [
            {
                "id": "technical-services",
                "name": "Technical Services Proposal",
                "category": "Technical Services",
                "description": "Standard template for technical consulting and services",
                "sections": [
                    "Executive Summary",
                    "Understanding of Requirements",
                    "Proposed Solution",
                    "Technical Approach",
                    "Project Timeline",
                    "Team and Resources",
                    "Risk Management",
                    "Budget and Investment",
                    "Terms and Conditions"
                ],
                "preview": "Comprehensive technical services proposal template with 9 sections."
            },
            {
                "id": "consulting",
                "name": "Management Consulting",
                "category": "Consulting",
                "description": "Template for management and strategy consulting proposals",
                "sections": [
                    "Executive Summary",
                    "Current State Analysis",
                    "Recommended Strategy",
                    "Implementation Roadmap",
                    "Change Management",
                    "Success Metrics",
                    "Our Expertise",
                    "Investment Required"
                ],
                "preview": "Strategic consulting template focusing on analysis, strategy, and implementation."
            },
            {
                "id": "software-dev",
                "name": "Software Development",
                "category": "Software Development",
                "description": "Template for custom software development projects",
                "sections": [
                    "Project Overview",
                    "Functional Requirements",
                    "Technical Architecture",
                    "Development Methodology",
                    "Quality Assurance",
                    "Deployment Strategy",
                    "Maintenance and Support",
                    "Project Timeline",
                    "Cost Breakdown"
                ],
                "preview": "Complete software development proposal covering architecture, methodology, QA, and deployment."
            },
            {
                "id": "research",
                "name": "Research and Analysis",
                "category": "Research",
                "description": "Template for research and market analysis projects",
                "sections": [
                    "Research Objectives",
                    "Methodology",
                    "Data Collection Plan",
                    "Analysis Framework",
                    "Deliverables",
                    "Timeline",
                    "Research Team",
                    "Budget"
                ],
                "preview": "Research-focused template with methodology, data collection, and analysis framework."
            }
        ]

        # Combine custom and predefined templates
        all_templates = custom_templates + predefined_templates
        print(f"📋 Total templates available: {len(all_templates)}")

        return {
            "status": "success",
            "templates": all_templates,
            "count": len(all_templates)
        }

    except Exception as e:
        print(f"❌ Error getting TOC templates: {e}")
        # Return predefined templates as fallback
        predefined_templates = [
            {
                "id": "technical-services",
                "name": "Technical Services Proposal",
                "category": "Technical Services",
                "description": "Standard template for technical consulting and services",
                "sections": [
                    "Executive Summary",
                    "Understanding of Requirements",
                    "Proposed Solution",
                    "Technical Approach",
                    "Project Timeline",
                    "Team and Resources",
                    "Risk Management",
                    "Budget and Investment",
                    "Terms and Conditions"
                ],
                "preview": "Comprehensive technical services proposal template with 9 sections."
            }
        ]
        return {
            "status": "success",
            "templates": predefined_templates,
            "count": len(predefined_templates)
        }

@app.delete("/delete_template/{template_id}")
async def delete_template(template_id: str):
    """
    Delete a template and all its associated data/files.
    """
    from toc_extractor import delete_template_by_id

    try:
        print(f"🗑️ Deleting template: {template_id}")
        result = delete_template_by_id(template_id)

        if result["status"] == "success":
            print(f"✅ Template {template_id} deleted successfully")
        else:
            print(f"❌ Failed to delete template {template_id}: {result['message']}")

        return result

    except Exception as e:
        print(f"❌ Error deleting template {template_id}: {e}")
        return {
            "status": "error",
            "message": str(e)
        }

@app.post("/apply_toc_template")
async def apply_toc_template(request: ApplyTOCTemplateRequest):
    """
    Apply a TOC template to create proposal sections.
    """
    from toc_extractor import apply_toc_template

    try:
        print(f"📋 Applying TOC template: {request.template_id}")

        sections = apply_toc_template(request.template_id, request.proposal_title)

        if not sections:
            return {
                "status": "error",
                "message": "Template not found or contains no sections"
            }

        print(f"✅ Created {len(sections)} sections from template")
        return {
            "status": "success",
            "sections": sections,
            "count": len(sections)
        }

    except Exception as e:
        print(f"❌ Error applying TOC template: {e}")
        return {
            "status": "error",
            "message": str(e)
        }

@app.get("/get_toc_preview/{template_id}")
async def get_toc_preview(template_id: str):
    """
    Get a markdown preview of a TOC template.
    """
    from toc_extractor import generate_toc_preview

    try:
        preview = generate_toc_preview(template_id)

        return {
            "status": "success",
            "preview": preview,
            "template_id": template_id
        }

    except Exception as e:
        print(f"❌ Error generating TOC preview: {e}")
        return {
            "status": "error",
            "message": str(e)
        }





