# doc_composer.py
from pathlib import Path
from typing import Any, Dict, List, Union
import re

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------- helpers ----------
def _normalize_text(content: Any) -> str:
    if content is None:
        return ""
    if isinstance(content, str):
        return content.strip()
    if isinstance(content, list):
        return "\n".join(_normalize_text(x) for x in content)
    if isinstance(content, dict):
        for k in ("content", "text", "body"):
            if k in content:
                return _normalize_text(content[k])
        return "\n".join(f"{k}: {_normalize_text(v)}" for k, v in content.items())
    return str(content)


def _slug(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", s.lower())


def _init_doc() -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    return doc


def _title_page(doc: Document, title: str = "Proposal Document", subtitle: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(12)

    doc.add_page_break()


def _insert_toc(doc: Document, levels: str = "1-3"):
    # Adds a TOC field. In Word, press F9 to build/update it.
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f'TOC \\o "{levels}" \\h \\z \\u')
    p._p.append(fld)
    doc.add_page_break()


def _set_heading(paragraph, level: int):
    paragraph.style = f"Heading {max(1, min(9, level))}"
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(3)


def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    _add_inline_runs(p, text)


def _add_body(doc: Document, text: str):
    if not text.strip():
        doc.add_paragraph("")  # blank line
        return
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    _add_inline_runs(p, text)


# very small inline markdown: **bold**, *italic*
_INLINE_TOK = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")

def _add_inline_runs(paragraph, text: str):
    pos = 0
    for m in _INLINE_TOK.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------- markdown renderer ----------
_HDR = re.compile(r"^(#{1,6})\s+(.*)$")
_BUL = re.compile(r"^(\s*)([-*])\s+(.*)$")

def _render_markdown_block(doc: Document, line: str, base_level: int):
    # Headings
    m = _HDR.match(line)
    if m:
        hashes, title = m.groups()
        md_level = len(hashes)  # 1..6
        level = max(1, min(9, base_level - 1 + md_level))  # relative to section level
        p = doc.add_paragraph(title.strip())
        _set_heading(p, level)
        return

    # Bullets (single level; nested bullets will still render as bullets but without outline levels)
    b = _BUL.match(line)
    if b:
        indent, _sym, txt = b.groups()
        _add_bullet(doc, txt.strip())
        return

    # Plain paragraph
    _add_body(doc, line)


def _strip_leading_duplicate_heading(lines: List[str], section_title: str) -> List[str]:
    # If the very first non-empty line is a markdown heading equal to the section title, drop it.
    for idx, raw in enumerate(lines):
        line = raw.strip()
        if not line:
            continue
        m = _HDR.match(line)
        if m:
            _, title = m.groups()
            if _slug(title) == _slug(section_title):
                return lines[idx + 1 :]  # drop the duplicate heading
        break
    return lines


# ---------- public API ----------
def compose_proposal_docx(
    sections: List[Dict[str, Any]],
    output_path: Union[str, Path],
    *,
    add_title_page: bool = True,
    add_toc: bool = True,
) -> None:
    """
    sections: [{ "title": str, "content": str, "level": int }]
    """
    doc = _init_doc()

    if add_title_page:
        _title_page(doc, "Proposal Document")
    if add_toc:
        _insert_toc(doc, "1-4")

    last_idx = len(sections) - 1
    for i, sec in enumerate(sections):
        title = sec.get("title") or "Section"
        content = _normalize_text(sec.get("content", ""))
        level = int(sec.get("level", 1))

        # Section heading from TOC (real Word Heading style)
        hp = doc.add_paragraph(title)
        _set_heading(hp, level)

        # Render body with markdown awareness, relative to this section's level
        lines = content.splitlines()
        lines = _strip_leading_duplicate_heading(lines, title)
        for raw in lines:
            _render_markdown_block(doc, raw.rstrip(), base_level=level)

        if i != last_idx:
            doc.add_page_break()

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"💾 Saved proposal to: {out}")
